from http.client import NON_AUTHORITATIVE_INFORMATION

from pdf2docx import Converter
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE

def parse_pagenum(s):
    s = s.replace('，', ',')
    s = s.replace('—', '-')
    group = s.split(',')
    temp = set()
    for each in group:
        if '-' in each:
            nums = [x for x in each.split('-') if x]
            if len(nums) == 2 and nums[0].isdigit() and nums[1].isdigit():
                temp.update(set(range(int(nums[0]), int(nums[1]) + 1)))
            else:
                return None
        elif each.isdigit():
            temp.add(int(each))
        else:
            return None
    output = sorted(list(temp))
    return output

def under_by(l, max):
    '''
    ``l`` is a list which include page numbers, start by 1. 
    ``max`` is the max page number.
    '''
    if max >= l[-1]:
        return l 
    length = len(l)
    if length <= 3:
        return [x for x in l if x <= max]
    p_start = 0
    p_end = length - 1
    while not (p_start + 1 == p_end):
        if l[(p_start + p_end) // 2] <= max:
            p_start = (p_start + p_end) // 2
        else:
            p_end = (p_start + p_end) // 2
    return l[:p_end]


def export_pagenum(l):
    temp_start = -1
    temp_end = -1
    output = list()
    for num in l:
        if temp_end + 1 == num:
            temp_end = num
        else:
            if temp_end != -1:
                if temp_end > temp_start:
                    output.append('{}-{}'.format(temp_start, temp_end))
                else:
                    output.append('{}'.format(temp_start))
            temp_start = num
            temp_end = num
    if temp_end > temp_start:
        output.append('{}-{}'.format(temp_start, temp_end))
    else:
        output.append('{}'.format(temp_start))

    return ','.join(output)

def set_text_style(docx_file):
    doc = Document(docx_file)
    saveDoc = Document()
    # 设置正文样式

    styles = doc.styles

    # title_style.font.bold = True  # 设置标题1的字体加粗
    # title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 设置标题1居中

    for section in doc.sections:
        # 设置页面顶部和底部边距
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    for style in styles:
        # if style.type == WD_STYLE_TYPE.PARAGRAPH:
        style.base_style = None

    title_style = styles.default(WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(14)  # 设置标题1的字体大小为14pt
    title_style.font.name = "SimSun"

    for i, paragraph in enumerate(doc.paragraphs):
        paragraph.style = title_style
        # 设置段落行距
        paragraph.paragraph_format.page_break_before = False
        paragraph.paragraph_format.line_spacing = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)  # 设置段前距
        paragraph.paragraph_format.space_after = Pt(0)  # 设置段后距
        paragraph.text = paragraph.text.strip().replace(" ", "")
        if len(paragraph.text) == 0:
            del doc.paragraphs[i]
            continue
        for j, run in enumerate(paragraph.runs):
            # run.text = run.text.strip()
            # run.text = run.text.rstrip("\r").replace(" ", "")

            paragraph.runs[j].font.spacing = Pt(0.5)  # 设置字间距为2pt
            paragraph.runs[j].font.name = "SimSun"  # 设置字体为新宋体
            paragraph.runs[j].font.size = Pt(16)  # 设置字体大小为12pt
            paragraph.runs[j].font.color.rgb = RGBColor(0, 0, 0)  # 设置字体颜色为红色

        saveDoc.add_paragraph(paragraph.text)

    # 设置文档最后一个段落为连续的页眉页脚
    doc.paragraphs[-1].paragraph_format.page_break_before = True
    doc.paragraphs[-1].paragraph_format.space_after = Pt(0)
        # 清除所有样式

    saveDoc.save(docx_file)


def convert_pdf_to_docx_with_style(pdf_file, docx_file):
    cv = Converter(pdf_file)
    cv.convert(docx_file, start=0, end=None)
    cv.close()

    set_text_style(docx_file)
